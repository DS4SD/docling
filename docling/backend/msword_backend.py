import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Any, Optional, Union

from docling_core.types.doc import (
    DocItemLabel,
    DoclingDocument,
    DocumentOrigin,
    GroupLabel,
    ImageRef,
    NodeItem,
    TableCell,
    TableData,
)
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tc
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree
from lxml.etree import XPath
from PIL import Image, UnidentifiedImageError
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument

_log = logging.getLogger(__name__)


class MsWordDocumentBackend(DeclarativeDocumentBackend):
    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.microsoft.com/office/word/2003/wordml"
        }
        # self.initialise(path_or_stream)
        # Word file:
        self.path_or_stream: Union[BytesIO, Path] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10
        self.level_at_new_list: Optional[int] = None
        self.parents: dict[int, Optional[NodeItem]] = {}
        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0

        self.history: dict[str, Any] = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }

        self.docx_obj = None
        try:
            if isinstance(self.path_or_stream, BytesIO):
                self.docx_obj = Document(self.path_or_stream)
            elif isinstance(self.path_or_stream, Path):
                self.docx_obj = Document(str(self.path_or_stream))

            self.valid = True
        except Exception as e:
            raise RuntimeError(
                f"MsPowerpointDocumentBackend could not load document with hash {self.document_hash}"
            ) from e

    @override
    def is_valid(self) -> bool:
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        return False

    @override
    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """Parses the DOCX into a structured document model.

        Returns:
            The parsed document.
        """

        origin = DocumentOrigin(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=self.document_hash,
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)
        if self.is_valid():
            assert self.docx_obj is not None
            doc = self.walk_linear(self.docx_obj.element.body, self.docx_obj, doc)
            return doc
        else:
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )

    def update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        self.history["names"].append(name)
        self.history["levels"].append(level)

        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def prev_name(self) -> Optional[str]:
        return self.history["names"][-1]

    def prev_level(self) -> Optional[int]:
        return self.history["levels"][-1]

    def prev_numid(self) -> Optional[int]:
        return self.history["numids"][-1]

    def prev_indent(self) -> Optional[int]:
        return self.history["indents"][-1]

    def get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v == None:
                return k
        return 0

    def walk_linear(
        self,
        body: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> DoclingDocument:
        for element in body:
            tag_name = etree.QName(element).localname
            # Check for Inline Images (blip elements)
            namespaces = {
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            }
            xpath_expr = XPath(".//a:blip", namespaces=namespaces)
            drawing_blip = xpath_expr(element)

            # Check for Tables
            if element.tag.endswith("tbl"):
                try:
                    self.handle_tables(element, docx_obj, doc)
                except Exception:
                    _log.debug("could not parse a table, broken docx table")

            elif drawing_blip:
                self.handle_pictures(docx_obj, drawing_blip, doc)
            # Check for the sdt containers, like table of contents
            elif tag_name in ["sdt"]:
                sdt_content = element.find(".//w:sdtContent", namespaces=namespaces)
                if sdt_content is not None:
                    # Iterate paragraphs, runs, or text inside <w:sdtContent>.
                    paragraphs = sdt_content.findall(".//w:p", namespaces=namespaces)
                    for p in paragraphs:
                        self.handle_text_elements(p, docx_obj, doc)
            # Check for Text
            elif tag_name in ["p"]:
                # "tcPr", "sectPr"
                self.handle_text_elements(element, docx_obj, doc)
            else:
                _log.debug(f"Ignoring element in DOCX with tag: {tag_name}")
        return doc

    def str_to_int(self, s: Optional[str], default: Optional[int] = 0) -> Optional[int]:
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def split_text_and_number(self, input_string: str) -> list[str]:
        match = re.match(r"(\D+)(\d+)$|^(\d+)(\D+)", input_string)
        if match:
            parts = list(filter(None, match.groups()))
            return parts
        else:
            return [input_string]

    def get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self.str_to_int(numId, None), self.str_to_int(ilvl, None)

        return None, None  # If the paragraph is not part of a list

    def get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        if paragraph.style is None:
            return "Normal", None
        label = paragraph.style.style_id
        if label is None:
            return "Normal", None
        if ":" in label:
            parts = label.split(":")

            if len(parts) == 2:
                return parts[0], self.str_to_int(parts[1], None)

        parts = self.split_text_and_number(label)

        if "Heading" in label and len(parts) == 2:
            parts.sort()
            label_str: str = ""
            label_level: Optional[int] = 0
            if parts[0] == "Heading":
                label_str = parts[0]
                label_level = self.str_to_int(parts[1], None)
            if parts[1] == "Heading":
                label_str = parts[1]
                label_level = self.str_to_int(parts[0], None)
            return label_str, label_level
        else:
            return label, None

    def handle_text_elements(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        paragraph = Paragraph(element, docx_obj)

        if paragraph.text is None:
            return
        text = paragraph.text.strip()

        # Common styles for bullet and numbered lists.
        # "List Bullet", "List Number", "List Paragraph"
        # Identify wether list is a numbered list or not
        # is_numbered = "List Bullet" not in paragraph.style.name
        is_numbered = False
        p_style_id, p_level = self.get_label_and_level(paragraph)
        numid, ilevel = self.get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # Handle lists
        if (
            numid is not None
            and ilevel is not None
            and p_style_id not in ["Title", "Heading"]
        ):
            self.add_listitem(
                doc,
                numid,
                ilevel,
                text,
                is_numbered,
            )
            self.update_history(p_style_id, p_level, numid, ilevel)
            return
        elif (
            numid is None
            and self.prev_numid() is not None
            and p_style_id not in ["Title", "Heading"]
        ):  # Close list
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0

        if p_style_id in ["Title"]:
            for key in range(len(self.parents)):
                self.parents[key] = None
            self.parents[0] = doc.add_text(
                parent=None, label=DocItemLabel.TITLE, text=text
            )
        elif "Heading" in p_style_id:
            self.add_header(doc, p_level, text)

        elif p_style_id in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "DefaultText",
            "ListParagraph",
            "ListBullet",
            "Quote",
        ]:
            level = self.get_level()
            doc.add_text(
                label=DocItemLabel.PARAGRAPH, parent=self.parents[level - 1], text=text
            )

        else:
            # Text style names can, and will have, not only default values but user values too
            # hence we treat all other labels as pure text
            level = self.get_level()
            doc.add_text(
                label=DocItemLabel.PARAGRAPH, parent=self.parents[level - 1], text=text
            )

        self.update_history(p_style_id, p_level, numid, ilevel)
        return

    def add_header(
        self, doc: DoclingDocument, curr_level: Optional[int], text: str
    ) -> None:
        level = self.get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # add invisible group
                for i in range(level, curr_level):
                    self.parents[i] = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
            elif curr_level < level:
                # remove the tail
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            self.parents[curr_level] = doc.add_heading(
                parent=self.parents[curr_level - 1],
                text=text,
                level=curr_level,
            )
        else:
            self.parents[self.level] = doc.add_heading(
                parent=self.parents[self.level - 1],
                text=text,
                level=1,
            )
        return

    def add_listitem(
        self,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        text: str,
        is_numbered: bool = False,
    ) -> None:
        enum_marker = ""

        level = self.get_level()
        prev_indent = self.prev_indent()
        if self.prev_numid() is None:  # Open new list
            self.level_at_new_list = level

            self.parents[level] = doc.add_group(
                label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
            )

            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[level],
                text=text,
            )

        elif (
            self.prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # Open indented list
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                # Determine if this is an unordered list or an ordered list.
                # Set GroupLabel.ORDERED_LIST when it fits.
                self.listIter = 0
                if is_numbered:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.ORDERED_LIST,
                        name="list",
                        parent=self.parents[i - 1],
                    )
                else:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[i - 1]
                    )

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[self.level_at_new_list + ilevel],
                text=text,
            )

        elif (
            self.prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # Close list
            for k, v in self.parents.items():
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[self.level_at_new_list + ilevel],
                text=text,
            )
            self.listIter = 0

        elif self.prev_numid() == numid or prev_indent == ilevel:
            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            doc.add_list_item(
                marker=enum_marker,
                enumerated=is_numbered,
                parent=self.parents[level - 1],
                text=text,
            )
        return

    def handle_tables(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        table: Table = Table(element, docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        _log.debug(f"Table grid with {num_rows} rows and {num_cols} columns")

        if num_rows == 1 and num_cols == 1:
            cell_element = table.rows[0].cells[0]
            # In case we have a table of only 1 cell, we consider it furniture
            # And proceed processing the content of the cell as though it's in the document body
            self.walk_linear(cell_element._element, docx_obj, doc)
            return

        data = TableData(num_rows=num_rows, num_cols=num_cols)
        cell_set: set[CT_Tc] = set()
        for row_idx, row in enumerate(table.rows):
            _log.debug(f"Row index {row_idx} with {len(row.cells)} populated cells")
            col_idx = 0
            while col_idx < num_cols:
                cell: _Cell = row.cells[col_idx]
                _log.debug(
                    f" col {col_idx} grid_span {cell.grid_span} grid_cols_before {row.grid_cols_before}"
                )
                if cell is None or cell._tc in cell_set:
                    _log.debug(f"  skipped since repeated content")
                    col_idx += cell.grid_span
                    continue
                else:
                    cell_set.add(cell._tc)

                spanned_idx = row_idx
                spanned_tc: Optional[CT_Tc] = cell._tc
                while spanned_tc == cell._tc:
                    spanned_idx += 1
                    spanned_tc = (
                        table.rows[spanned_idx].cells[col_idx]._tc
                        if spanned_idx < num_rows
                        else None
                    )
                _log.debug(f"  spanned before row {spanned_idx}")

                table_cell = TableCell(
                    text=cell.text,
                    row_span=spanned_idx - row_idx,
                    col_span=cell.grid_span,
                    start_row_offset_idx=row.grid_cols_before + row_idx,
                    end_row_offset_idx=row.grid_cols_before + spanned_idx,
                    start_col_offset_idx=col_idx,
                    end_col_offset_idx=col_idx + cell.grid_span,
                    col_header=False,
                    row_header=False,
                )
                data.table_cells.append(table_cell)
                col_idx += cell.grid_span

        level = self.get_level()
        doc.add_table(data=data, parent=self.parents[level - 1])
        return

    def handle_pictures(
        self, docx_obj: DocxDocument, drawing_blip: Any, doc: DoclingDocument
    ) -> None:
        def get_docx_image(drawing_blip):
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId in docx_obj.part.rels:
                # Access the image part using the relationship ID
                image_part = docx_obj.part.rels[rId].target_part
                image_data = image_part.blob  # Get the binary image data
            return image_data

        level = self.get_level()
        # Open the BytesIO object with PIL to create an Image
        try:
            image_data = get_docx_image(drawing_blip)
            image_bytes = BytesIO(image_data)
            pil_image = Image.open(image_bytes)
            doc.add_picture(
                parent=self.parents[level - 1],
                image=ImageRef.from_pil(image=pil_image, dpi=72),
                caption=None,
            )
        except (UnidentifiedImageError, OSError) as e:
            _log.warning("Warning: image cannot be loaded by Pillow")
            doc.add_picture(
                parent=self.parents[level - 1],
                caption=None,
            )
        return
